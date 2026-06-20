"""
Ejecutar una sola vez para crear la plantilla base:
    python crear_plantilla.py
"""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Cm
from pathlib import Path

DESTINO = Path("app/templates/certificado_template.docx")
DESTINO.parent.mkdir(parents=True, exist_ok=True)


def _titulo(doc: Document, texto: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(texto)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x1A, 0x53, 0x76)


def _subtitulo(doc: Document, texto: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(texto)
    run.bold = True
    run.font.size = Pt(13)


def _linea(doc: Document, etiqueta: str, placeholder: str) -> None:
    p = doc.add_paragraph()
    run_label = p.add_run(f"{etiqueta}: ")
    run_label.bold = True
    run_label.font.size = Pt(11)
    run_valor = p.add_run(placeholder)
    run_valor.font.size = Pt(11)


def _espacio(doc: Document) -> None:
    doc.add_paragraph()


doc = Document()

# Márgenes
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(3)

_titulo(doc, "CERTIFICADO DE RECOLECCIÓN DE ACEITE USADO")
_espacio(doc)
_subtitulo(doc, "Código: {{codigo}}")
_espacio(doc)

_linea(doc, "Restaurante",         "{{restaurante}}")
_linea(doc, "NIT",                 "{{nit}}")
_linea(doc, "Dirección",           "{{direccion}}")
_linea(doc, "Ciudad",              "{{ciudad}}")
_espacio(doc)
_linea(doc, "Fecha de recolección","{{fecha_recoleccion}}")
_linea(doc, "Cantidad",            "{{cantidad}}")
_linea(doc, "Tipo",                "{{tipo}}")
_espacio(doc)
_linea(doc, "Fecha de generación", "{{fecha_generacion}}")
_espacio(doc)

p_firma = doc.add_paragraph()
p_firma.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_firma.add_run("_" * 40)

p_cargo = doc.add_paragraph()
p_cargo.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p_cargo.add_run("Responsable de Recolección")
run.font.size = Pt(10)

doc.save(DESTINO)
print(f"Plantilla creada en: {DESTINO}")
