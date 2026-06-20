"""
Copia la plantilla original y reemplaza los valores reales con placeholders.
Ejecutar una sola vez:
    python preparar_plantilla.py
"""
import shutil
from pathlib import Path
from docx import Document

ORIGEN  = Path(r"C:\Users\santi\Downloads\CERTIFICADO DE RECOLECCIÓN Y APROVECHAMIENTO DE ACU BIO A GRASS.docx - ABRIL_2026 (2).docx")
DESTINO = Path("app/templates/certificado_template.docx")

DESTINO.parent.mkdir(parents=True, exist_ok=True)
shutil.copy2(ORIGEN, DESTINO)

doc = Document(DESTINO)


def reemplazar_celda(tabla_idx: int, fila_idx: int, celda_idx: int, nuevo_texto: str) -> None:
    celda = doc.tables[tabla_idx].rows[fila_idx].cells[celda_idx]
    for p in celda.paragraphs:
        if p.runs:
            p.runs[0].text = nuevo_texto
            for r in p.runs[1:]:
                r.text = ""


# ── Tabla 0: Información del generador ───────────────────────────────────────
reemplazar_celda(0, 2, 0, "{{restaurante}}")
reemplazar_celda(0, 2, 1, "{{nit}}")
reemplazar_celda(0, 2, 2, "{{direccion}}")
reemplazar_celda(0, 2, 3, "{{telefono}}")     # campo futuro; queda vacío si no se pasa

# ── Tabla 1: Información de recolección ──────────────────────────────────────
reemplazar_celda(1, 2, 0, "{{fecha_recoleccion}}")
reemplazar_celda(1, 2, 1, "{{descripcion_tipo}}")
reemplazar_celda(1, 2, 2, "{{cantidad}}")
reemplazar_celda(1, 2, 3, "{{descuento}}")
reemplazar_celda(1, 2, 4, "{{total}}")

# ── Tabla 2: Aprovechamiento ──────────────────────────────────────────────────
reemplazar_celda(2, 2, 1, "{{total}}")

# ── Párrafo de firma (índice 11) ──────────────────────────────────────────────
p_firma = doc.paragraphs[11]
texto_original = p_firma.text
if p_firma.runs:
    p_firma.runs[0].text = (
        "Para mayor constancia se firma en {{ciudad}}, el día {{fecha_generacion}}."
    )
    for r in p_firma.runs[1:]:
        r.text = ""

# ── Código de certificado (párrafo 1, vacío bajo el título) ──────────────────
p_codigo = doc.paragraphs[1]
if p_codigo.runs:
    p_codigo.runs[0].text = "Código: {{codigo}}"
else:
    run = p_codigo.add_run("Código: {{codigo}}")

doc.save(DESTINO)
print(f"Plantilla preparada en: {DESTINO}")
print("Placeholders insertados:")
print("  {{restaurante}}, {{nit}}, {{direccion}}, {{telefono}}")
print("  {{fecha_recoleccion}}, {{descripcion_tipo}}, {{cantidad}}, {{descuento}}, {{total}}")
print("  {{ciudad}}, {{fecha_generacion}}, {{codigo}}")
