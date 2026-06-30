import re
import shutil
import subprocess
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt

from app.config import settings
from app.models.certificate import Certificate, TipoCertificado
from app.utils.file_utils import ruta_documento
from app.utils.logger import get_logger

logger = get_logger(__name__)

_SOFFICE_PATHS = [
    Path("C:/Program Files/LibreOffice/program/soffice.exe"),          # Windows
    Path("C:/Program Files (x86)/LibreOffice/program/soffice.exe"),    # Windows 32-bit
    Path("/Applications/LibreOffice.app/Contents/MacOS/soffice"),      # Mac
    Path("/usr/lib/libreoffice/program/soffice"),                       # Linux
]
_TEMPLATE_FILENAMES = {
    1: "certificado_template.docx",
    2: "certificado_template_2.docx",
    3: "certificado_template_3.docx",
}
_TEMPLATES = {
    template_id: settings.templates_dir / filename
    for template_id, filename in _TEMPLATE_FILENAMES.items()
}
_PLACEHOLDER_PATTERN = re.compile(r"\{\{[^}]+\}\}")

_DESCRIPCION_TIPO = {
    TipoCertificado.PIMPINA.value: "Aceite de cocina usado – ACU (Pimpinas de 20 litros)",
    TipoCertificado.KG.value:      "Aceite de cocina usado – ACU",
}


def get_available_templates() -> dict[int, str]:
    available: dict[int, str] = {}
    for template_id, filename in _TEMPLATE_FILENAMES.items():
        template_path = settings.templates_dir / filename
        if template_path.exists() and template_path.is_file():
            available[template_id] = filename
    return available


def _find_soffice() -> Path | None:
    for path in _SOFFICE_PATHS:
        if path.exists():
            return path
    found = shutil.which("soffice")
    return Path(found) if found else None


def _get_template_path(plantilla: int) -> Path:
    template_path = _TEMPLATES.get(plantilla)
    if template_path is None:
        raise ValueError(f"Plantilla {plantilla} no soportada. Usa {list(_TEMPLATES)}.")
    return template_path


def _validar_template_path(template_path: Path) -> None:
    if not template_path.exists():
        raise FileNotFoundError(f"Plantilla no encontrada: {template_path}")
    if not template_path.is_file():
        raise FileNotFoundError(f"La plantilla debe ser un archivo: {template_path}")
    if template_path.suffix.lower() != ".docx":
        raise ValueError(f"Plantilla inválida: {template_path.name} debe tener extensión .docx")


def _scan_template_for_placeholders(template_path: Path) -> set[str]:
    """Lee placeholders consolidando runs para no perderse texto partido entre runs."""
    found: set[str] = set()
    try:
        doc = Document(template_path)
    except Exception as exc:
        raise ValueError(f"La plantilla {template_path.name} no es un archivo DOCX válido.") from exc

    def _scan_paragraphs(paragraphs) -> None:
        for p in paragraphs:
            texto = "".join(r.text for r in p.runs)
            found.update(_PLACEHOLDER_PATTERN.findall(texto))

    _scan_paragraphs(doc.paragraphs)
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                _scan_paragraphs(celda.paragraphs)
    for seccion in doc.sections:
        _scan_paragraphs(seccion.header.paragraphs)
        _scan_paragraphs(seccion.footer.paragraphs)

    return found


def _verify_template_has_placeholders(template_path: Path) -> set[str]:
    return _scan_template_for_placeholders(template_path)


def _insert_codigo_en_header(doc: Document, codigo: str) -> None:
    def _write_text(header):
        paragraph = header.add_paragraph()
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(codigo)
        font = run.font
        font.size = Pt(10)
        font.bold = True
        font.name = "Arial"

    for section in doc.sections:
        _write_text(section.header)
        if section.first_page_header is not section.header:
            _write_text(section.first_page_header)

    if doc.paragraphs:
        first_par = doc.paragraphs[0].insert_paragraph_before()
        first_par.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        first_par.paragraph_format.space_before = Pt(0)
        first_par.paragraph_format.space_after = Pt(0)
        run = first_par.add_run(codigo)
        font = run.font
        font.size = Pt(10)
        font.bold = True
        font.name = "Arial"


def _reemplazar_parrafo(parrafo, replacements: dict[str, str]) -> None:
    # Primero reemplaza dentro de cada run individual (preserva formato)
    for run in parrafo.runs:
        for k, v in replacements.items():
            if k in run.text:
                run.text = run.text.replace(k, v)

    # Si quedó algún placeholder partido entre runs, consolida solo ese párrafo
    texto = "".join(r.text for r in parrafo.runs)
    if not any(k in texto for k in replacements):
        return
    for k, v in replacements.items():
        texto = texto.replace(k, v)
    if parrafo.runs:
        parrafo.runs[0].text = texto
        for r in parrafo.runs[1:]:
            r.text = ""


def _reemplazar_documento(doc: Document, replacements: dict[str, str]) -> None:
    for p in doc.paragraphs:
        _reemplazar_parrafo(p, replacements)

    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for p in celda.paragraphs:
                    _reemplazar_parrafo(p, replacements)

    for seccion in doc.sections:
        for p in seccion.header.paragraphs:
            _reemplazar_parrafo(p, replacements)
        for p in seccion.footer.paragraphs:
            _reemplazar_parrafo(p, replacements)


def _placeholders_left_in_doc(doc: Document) -> list[str]:
    remaining: set[str] = set()

    def _scan_paragraphs(paragraphs: Iterable) -> None:
        for par in paragraphs:
            remaining.update(_PLACEHOLDER_PATTERN.findall(par.text))

    _scan_paragraphs(doc.paragraphs)
    for tabla in doc.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                _scan_paragraphs(celda.paragraphs)

    for seccion in doc.sections:
        _scan_paragraphs(seccion.header.paragraphs)
        _scan_paragraphs(seccion.footer.paragraphs)

    return sorted(remaining)


def _convertir_a_pdf(ruta_docx: Path) -> Path:
    """
    Convierte un .docx a .pdf usando LibreOffice headless.
    Devuelve la ruta del PDF generado.
    """
    soffice = _find_soffice()
    if not soffice:
        raise FileNotFoundError(
            "LibreOffice no encontrado. Instala LibreOffice o agrega 'soffice' al PATH."
        )

    resultado = subprocess.run(
        [
            str(soffice),
            "--headless",
            "--convert-to", "pdf",
            "--outdir", str(ruta_docx.parent),
            str(ruta_docx),
        ],
        capture_output=True,
        text=True,
        timeout=60,
    )

    ruta_pdf = ruta_docx.with_suffix(".pdf")

    if resultado.returncode != 0 or not ruta_pdf.exists():
        raise RuntimeError(
            f"Error al convertir a PDF: {resultado.stderr or resultado.stdout}"
        )

    logger.info("PDF generado: %s", ruta_pdf.name)
    return ruta_pdf


def generar_certificado_docx(certificado: Certificate, plantilla: int = 1) -> Path:
    """
    Carga la plantilla indicada, valida placeholders, reemplaza los valores
    y guarda el documento generado.

    Si LibreOffice está disponible, también intenta convertir el resultado a PDF.
    """
    template_path = _get_template_path(plantilla)
    _validar_template_path(template_path)
    placeholders = _verify_template_has_placeholders(template_path)

    cantidad_str = str(certificado.cantidad)
    descripcion = _DESCRIPCION_TIPO.get(
        certificado.tipo,
        "Aceite de cocina usado – ACU",
    )

    # Mapa de todos los posibles placeholders (con y sin tildes/espacios)
    _VALORES: dict[str, str] = {
        "{{codigo}}":              certificado.codigo_certificado,
        "{{restaurante}}":         certificado.restaurante,
        "{{nit}}":                 certificado.nit,
        "{{direccion}}":           certificado.direccion,
        "{{dirección}}":           certificado.direccion,
        "{{telefono}}":            "",
        "{{teléfono}}":            "",
        "{{ciudad}}":              certificado.ciudad,
        "{{fecha_recoleccion}}":   certificado.fecha_recoleccion.strftime("%d/%m/%Y"),
        "{{fecha recolección}}":   certificado.fecha_recoleccion.strftime("%d/%m/%Y"),
        "{{descripcion_tipo}}":    descripcion,
        "{{descripción_tipo}}":    descripcion,
        "{{cantidad}}":            cantidad_str,
        "{{descuento}}":           "0",
        "{{total}}":               cantidad_str,
        "{{fecha_generacion}}":    certificado.fecha_generacion.strftime("%d/%m/%Y"),
        "{{fecha generación}}":    certificado.fecha_generacion.strftime("%d/%m/%Y"),
        "{{dia}}":                 str(certificado.fecha_recoleccion.day),
    }
    # Siempre incluir todos — _reemplazar_parrafo consolida runs y atrapa los partidos
    replacements = _VALORES

    doc = Document(template_path)
    _reemplazar_documento(doc, replacements)

    ruta_docx = ruta_documento(
        restaurante=certificado.restaurante,
        fecha=certificado.fecha_recoleccion,
        codigo=certificado.codigo_certificado,
    )
    doc.save(ruta_docx)
    logger.info("DOCX generado: %s", ruta_docx)

    try:
        return _convertir_a_pdf(ruta_docx)
    except FileNotFoundError as exc:
        logger.warning("LibreOffice no disponible: %s. Se conservará DOCX.", exc)
        return ruta_docx
    except RuntimeError as exc:
        logger.warning("Error de conversión a PDF, se conservará DOCX: %s", exc)
        return ruta_docx
